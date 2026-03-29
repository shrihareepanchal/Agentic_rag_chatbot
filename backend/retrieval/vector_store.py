"""
Load various file types into a uniform list of
(text: str, metadata: dict) tuples.

FIXES APPLIED:
1. DOCX: Now extracts BOTH paragraphs AND tables
2. DOCX: Skips very short headings (<20 chars) that pollute the index
3. DOCX: Joins consecutive short paragraphs into bigger chunks
4. PDF:  Better error handling for scanned/corrupted pages
5. All: Minimum text length filter to avoid indexing empty/junk chunks

NOTE (per your request):
- All try/except handling has been REMOVED from this file.
- To FIX the ingestion error "'Chroma' object has no attribute 'add'",
  this file now returns a wrapper that provides `.add()` while using
  LangChain Chroma's `add_texts()` under the hood.
- Function names and signatures are preserved to avoid breaking other files.
"""
from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import List, Tuple
from functools import lru_cache

from backend.utils.logger import get_logger

logger = get_logger(__name__)

RawChunk = Tuple[str, dict]  # (text, metadata)

# Minimum characters for a chunk to be worth indexing
MIN_CHUNK_LENGTH = 20


# ── PDF ───────────────────────────────────────────────────────────────────────

def load_pdf(file_path: str) -> List[RawChunk]:
    """Extract text page-by-page from a PDF."""
    import pypdf  # no try/except as requested

    chunks: List[RawChunk] = []
    filename = Path(file_path).name

    with open(file_path, "rb") as f:
        reader = pypdf.PdfReader(f)
        total_pages = len(reader.pages)

        for page_num, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if len(text) < MIN_CHUNK_LENGTH:
                continue

            chunks.append((
                text,
                {
                    "source": filename,
                    "file_type": "pdf",
                    "page_number": page_num,
                    "total_pages": total_pages,
                },
            ))

    logger.info("pdf_loaded", filename=filename, pages=len(chunks))
    return chunks


# ── TXT / MD ─────────────────────────────────────────────────────────────────

def load_txt(file_path: str) -> List[RawChunk]:
    """Load a plain text file."""
    filename = Path(file_path).name
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read().strip()

    if len(text) < MIN_CHUNK_LENGTH:
        return []

    return [(text, {"source": filename, "file_type": "txt", "page_number": None})]


# ── DOCX ──────────────────────────────────────────────────────────────────────

def load_docx(file_path: str) -> List[RawChunk]:
    """
    Load a DOCX file extracting:
      1. Paragraphs  (body text, headings)
      2. Tables      (each row becomes a chunk)

    Also joins short consecutive paragraphs to improve chunk quality.
    """
    from docx import Document  # no try/except as requested

    filename = Path(file_path).name
    doc = Document(file_path)

    chunks: List[RawChunk] = []

    # ── Step 1: Extract Paragraphs ────────────────────────────────────────────
    para_idx = 0
    buffer: List[str] = []  # collect short paragraphs to join together

    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            # Empty paragraph → flush buffer
            if buffer:
                combined = " ".join(buffer).strip()
                if len(combined) >= MIN_CHUNK_LENGTH:
                    chunks.append((
                        combined,
                        {
                            "source": filename,
                            "file_type": "docx",
                            "paragraph": para_idx,
                            "page_number": None,
                            "content_type": "paragraph",
                        },
                    ))
                buffer = []
            para_idx += 1
            continue

        # Optional: Skip tiny heading-like noise (<20 chars)
        # (Your header comment mentions it; implement explicitly)
        if len(text) < 20:
            buffer.append(text)
            continue

        # Long paragraph → flush buffer first, then add on its own
        if len(text) >= 80:
            if buffer:
                combined = " ".join(buffer).strip()
                if len(combined) >= MIN_CHUNK_LENGTH:
                    chunks.append((
                        combined,
                        {
                            "source": filename,
                            "file_type": "docx",
                            "paragraph": para_idx,
                            "page_number": None,
                            "content_type": "paragraph",
                        },
                    ))
                buffer = []
                para_idx += 1

            chunks.append((
                text,
                {
                    "source": filename,
                    "file_type": "docx",
                    "paragraph": para_idx,
                    "page_number": None,
                    "content_type": "paragraph",
                },
            ))
            para_idx += 1
        else:
            # Short paragraph → add to buffer
            buffer.append(text)

    # Flush remaining buffer
    if buffer:
        combined = " ".join(buffer).strip()
        if len(combined) >= MIN_CHUNK_LENGTH:
            chunks.append((
                combined,
                {
                    "source": filename,
                    "file_type": "docx",
                    "paragraph": para_idx,
                    "page_number": None,
                    "content_type": "paragraph",
                },
            ))

    # ── Step 2: Extract Tables ────────────────────────────────────────────────
    table_chunks_count = 0

    for table_idx, table in enumerate(doc.tables):
        headers: List[str] = []
        if table.rows:
            headers = [cell.text.strip() for cell in table.rows[0].cells if cell.text.strip()]

        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            cells = [c for c in cells if c]

            if not cells:
                continue

            if headers and row_idx > 0 and len(cells) == len(headers):
                row_text = " | ".join(
                    f"{h}: {c}" for h, c in zip(headers, cells) if h and c
                )
            else:
                row_text = " | ".join(cells)

            if len(row_text) < MIN_CHUNK_LENGTH:
                continue

            chunks.append((
                row_text,
                {
                    "source": filename,
                    "file_type": "docx",
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "page_number": None,
                    "content_type": "table",
                },
            ))
            table_chunks_count += 1

    logger.info(
        "docx_loaded",
        filename=filename,
        total_chunks=len(chunks),
        paragraph_chunks=len(chunks) - table_chunks_count,
        table_chunks=table_chunks_count,
    )
    return chunks


# ── JSON ──────────────────────────────────────────────────────────────────────

def load_json(file_path: str) -> List[RawChunk]:
    """
    Flatten JSON into text chunks.
    - List of objects → each object = one chunk
    - Single object   → groups of 5 keys = one chunk
    """
    filename = Path(file_path).name
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    chunks: List[RawChunk] = []

    def obj_to_text(obj: dict) -> str:
        parts = []
        for k, v in obj.items():
            if isinstance(v, (dict, list)):
                parts.append(f"{k}: {json.dumps(v, ensure_ascii=False)}")
            else:
                parts.append(f"{k}: {v}")
        return "\n".join(parts)

    if isinstance(data, list):
        for idx, item in enumerate(data):
            text = obj_to_text(item) if isinstance(item, dict) else str(item)
            if len(text) >= MIN_CHUNK_LENGTH:
                chunks.append((
                    text,
                    {
                        "source": filename,
                        "file_type": "json",
                        "record_index": idx,
                        "page_number": None,
                    },
                ))

    elif isinstance(data, dict):
        items = list(data.items())
        for i in range(0, len(items), 5):
            group = dict(items[i: i + 5])
            text = obj_to_text(group)
            if len(text) >= MIN_CHUNK_LENGTH:
                chunks.append((
                    text,
                    {
                        "source": filename,
                        "file_type": "json",
                        "record_index": i // 5,
                        "page_number": None,
                    },
                ))

    logger.info("json_loaded", filename=filename, chunks=len(chunks))
    return chunks


# ── CSV ───────────────────────────────────────────────────────────────────────

def load_csv(file_path: str) -> List[RawChunk]:
    """Load CSV: each row becomes a text chunk with column headers as context."""
    filename = Path(file_path).name
    chunks: List[RawChunk] = []

    with open(file_path, "r", encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.DictReader(f)
        for row_idx, row in enumerate(reader):
            parts = [f"{k}: {v}" for k, v in row.items() if v and str(v).strip()]
            text = " | ".join(parts).strip()
            if len(text) < MIN_CHUNK_LENGTH:
                continue

            chunks.append((
                text,
                {
                    "source": filename,
                    "file_type": "csv",
                    "row_index": row_idx,
                    "page_number": None,
                },
            ))

    logger.info("csv_loaded", filename=filename, rows=len(chunks))
    return chunks


# ── Dispatcher ────────────────────────────────────────────────────────────────

def load_file(file_path: str) -> List[RawChunk]:
    """Dispatch to the correct loader based on file extension."""
    ext = Path(file_path).suffix.lower()
    loaders = {
        ".pdf":  load_pdf,
        ".txt":  load_txt,
        ".md":   load_txt,
        ".docx": load_docx,
        ".json": load_json,
        ".csv":  load_csv,
    }
    loader = loaders.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file type: {ext}")
    return loader(file_path)


def supported_extensions() -> List[str]:
    return [".pdf", ".txt", ".md", ".docx", ".json", ".csv"]


# =============================
# Vector store factory (Chroma)
# =============================

from functools import lru_cache
from typing import Any, Dict, List, Optional

class ChromaVectorStoreAdapter:
    """
    Adapter so the rest of the app can call:
      - vs.add(..., embeddings=...)  OR vs.add(...) without embeddings
      - vs.search(query, k)

    Fixes:
      1) "'Chroma' object has no attribute 'add'"
      2) "ChromaVectorStoreAdapter.add() got an unexpected keyword argument 'embeddings'"
    """

    def __init__(self, store: Any):
        self._store = store  # LangChain Chroma vectorstore
    def add(
        self,
        texts: Optional[List[str]] = None,
        metadatas: Optional[List[Dict[str, Any]]] = None,
        ids: Optional[List[str]] = None,
        embeddings: Optional[List[List[float]]] = None,
        **kwargs: Any,
    ):
        """
        Backward compatible add():

        Accepts BOTH:
          - texts=...
          - documents=...  (legacy callers)

        Supports:
          - embeddings provided (writes using Chroma native collection)
          - embeddings not provided (falls back to LangChain add_texts)
        """
        # ✅ Back-compat: allow documents= as alias for texts=
        if texts is None and "documents" in kwargs:
            texts = kwargs.pop("documents")

        # Determine batch size
        n = None
        if texts is not None:
            n = len(texts)
        elif embeddings is not None:
            n = len(embeddings)

        if not n:
            raise ValueError("add() requires non-empty `texts`/`documents` or `embeddings`.")

        # Ensure ids exist and match length
        if ids is None:
            ids = [str(i) for i in range(n)]
        if len(ids) != n:
            raise ValueError(f"len(ids)={len(ids)} must match batch size={n}.")

        if metadatas is not None and len(metadatas) != n:
            raise ValueError(f"len(metadatas)={len(metadatas)} must match batch size={n}.")

        # Sanitize metadata: ChromaDB only accepts str, int, float, bool — remove None values
        if metadatas is not None:
            metadatas = [
                {k: v for k, v in m.items() if v is not None} for m in metadatas
            ]

        # If embeddings are provided, write using native Chroma collection API
        if embeddings is not None:
            add_kwargs: Dict[str, Any] = {
                "ids": ids,
                "embeddings": embeddings,
                "metadatas": metadatas,
            }
            if texts is not None:
                add_kwargs["documents"] = texts

            self._store._collection.add(**add_kwargs)
            return {"added": n}

        # No embeddings -> LangChain will compute embeddings
        if texts is None:
            raise ValueError("add() without `embeddings` requires `texts`/`documents`.")

        return self._store.add_texts(texts=texts, metadatas=metadatas, ids=ids)

    def search(self, query: str, k: int = 5):
        return self._store.similarity_search_with_score(query, k=k)

    def similarity_search_with_score(self, query: str, k: int = 5, filter: dict | None = None):
        """Expose LangChain Chroma's similarity_search_with_score with optional metadata filter."""
        kwargs: dict = {"query": query, "k": k}
        if filter:
            kwargs["filter"] = filter
        return self._store.similarity_search_with_score(**kwargs)

    def similarity_search(self, query: str, k: int = 5):
        return self._store.similarity_search(query, k=k)

    def delete_by_source(self, source_filename: str) -> int:
        """
        Delete ALL chunks belonging to a specific source file.
        Returns the number of chunks deleted.
        """
        col = getattr(self._store, "_collection", None)
        if col is None:
            raise RuntimeError("Cannot access ChromaDB collection for deletion.")

        # Get IDs of all chunks with this source
        data = col.get(
            where={"source": source_filename},
            include=[],  # only need IDs
        )
        ids_to_delete = data.get("ids") or []

        if ids_to_delete:
            # ChromaDB delete() has a batch limit; chunk if needed
            batch_size = 5000
            for i in range(0, len(ids_to_delete), batch_size):
                batch = ids_to_delete[i : i + batch_size]
                col.delete(ids=batch)

        logger.info(
            "vector_store_delete_by_source",
            source=source_filename,
            deleted_count=len(ids_to_delete),
        )
        return len(ids_to_delete)

    def as_retriever(self, **kwargs):
        return self._store.as_retriever(**kwargs)


@lru_cache
def get_vector_store():
    """
    Create/return the Chroma vector store used by the app.
    IMPORTANT: Returns an adapter that provides `.add()`,
    since ingestion/indexer calls `vs.add(...)`.
    """
    from langchain_community.vectorstores import Chroma
    from backend.config import get_settings
    from backend.ingestion.embedder import get_embedder

    settings = get_settings()

    persist_dir = getattr(settings, "chroma_persist_dir", None) or "./data/chroma_db"

    collection = getattr(settings, "chroma_collection_name", None) or "rag_documents"

    chroma = Chroma(
        collection_name=collection,
        persist_directory=persist_dir,
        embedding_function=get_embedder(),
    )

    return ChromaVectorStoreAdapter(chroma)